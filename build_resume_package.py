#!/usr/bin/env python3
"""
build_resume_package.py
=======================
Builds Bilal Abdulkadir Muhammed's full Centennial College 2026 Residence
Scholarship application package as a single Word document.

Output structure (7 pages):
  1-2. Resume
  3.   References
  4-5. Personal Statement
  6-7. Leadership Essay

Usage:
    python3 build_resume_package.py
    python3 build_resume_package.py --output /path/to/output.docx

Dependencies:
    pip install python-docx
"""

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed. Run:  pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ============================================================================
# CONFIG
# ============================================================================
DEFAULT_OUTPUT = "/workspace/Bilal_Abdulkadir_Resume.docx"

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)   # deep navy
TEXT   = RGBColor(0x20, 0x20, 0x20)   # near-black
MUTED  = RGBColor(0x55, 0x55, 0x55)   # secondary text


# ============================================================================
# LOW-LEVEL OPENXML HELPERS
# ============================================================================
def _oxml(tag):
    return OxmlElement(tag)


def set_run_color(run, rgb):
    """Set a run's font color from an RGBColor."""
    run.font.color.rgb = rgb


def add_horizontal_rule(paragraph, color="1F3A5F", size="6"):
    """Add a bottom border under a paragraph (used for section headings + gold rules)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = _oxml("w:pBdr")
    bottom = _oxml("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_break(doc):
    """Insert a hard page break."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run()
    br = _oxml("w:br")
    br.set(qn("w:type"), "page")
    r._r.append(br)


# ============================================================================
# PARAGRAPH BUILDERS
# ============================================================================
def add_section_heading(doc, title):
    """UPPERCASE section heading with navy rule underneath."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    add_horizontal_rule(p, color="1F3A5F", size="8")
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Calibri"
    set_run_color(r, ACCENT)
    return p


def add_bullet(doc, text, bold_lead=None):
    """Bulleted list item; optional bold lead-in followed by normal tail text."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True
        r1.font.size = Pt(10.5)
        set_run_color(r1, TEXT)
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
        set_run_color(r2, TEXT)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        set_run_color(r, TEXT)
    return p


def add_job_title(doc, title, org_and_dates):
    """Job header line: bold navy title + muted italic org / dates."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.size = Pt(11)
    set_run_color(r1, ACCENT)
    r2 = p.add_run("    " + org_and_dates)
    r2.font.size = Pt(10)
    r2.italic = True
    set_run_color(r2, MUTED)


def add_centered_title(doc, text, size=22):
    """Large centered page title (used on References / Statement / Essay pages)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    set_run_color(r, ACCENT)
    return p


def add_centered_subtitle(doc, text):
    """Italic centered muted subtitle line under page title."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    set_run_color(r, MUTED)
    return p


def add_gold_rule(doc):
    """Decorative gold rule used under page titles."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_horizontal_rule(p, color="C9A227", size="12")


def add_essay_paragraph(doc, text, italic=False, bold=False):
    """Justified essay body paragraph with 1.25 line spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(11)
    set_run_color(r, TEXT)
    r.italic = italic
    r.bold = bold
    return p


def add_reference_card(doc, name, title, relationship, contact_line):
    """Single reference entry: bold navy name, italic title, plain relationship, contact line."""
    # Name
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(12)
    set_run_color(r, ACCENT)
    # Title
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(title)
    r.italic = True
    r.font.size = Pt(10.5)
    set_run_color(r, MUTED)
    # Relationship
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    r = p3.add_run(relationship)
    r.font.size = Pt(10.5)
    set_run_color(r, TEXT)
    # Contact
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(0)
    r = p4.add_run(contact_line)
    r.font.size = Pt(10)
    set_run_color(r, TEXT)


# ============================================================================
# PAGE BUILDERS
# ============================================================================
def build_resume(doc):
    """Pages 1-2: Resume."""

    # ---------- Title block ----------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    tr = title_p.add_run("Bilal Abdulkadir Muhammed")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.name = "Calibri"
    set_run_color(tr, ACCENT)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(2)
    sr = sub_p.add_run("Centennial College 2026 Residence Scholarship  ·  Incoming Student, Fall 2026")
    sr.italic = True
    sr.font.size = Pt(11.5)
    set_run_color(sr, MUTED)

    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_after = Pt(4)
    add_horizontal_rule(rule_p, color="C9A227", size="12")

    # ---------- Contact ----------
    for line in [
        "Addis Ababa, Ethiopia  ·  Open to Relocation to Canada",
        "+251 941 322 948  ·  bilalabdulkadir286@gmail.com",
        "linkedin.com/in/bilalabdulkadir  ·  github.com/Bilalabdulkadir  ·  bilalabdulkadir.github.io",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.size = Pt(10)
        set_run_color(r, MUTED)

    doc.add_paragraph()  # spacer

    # ---------- Professional Summary ----------
    add_section_heading(doc, "Professional Summary")
    add_essay_paragraph(doc,
        "Detail-oriented IT professional with 6+ years of experience in technical support and "
        "database management, transitioning into data analytics with strong proficiency in Python, "
        "Power BI, and cloud computing. As an incoming Centennial College student (Fall 2026), I am "
        "committed to academic excellence, peer mentorship, and contributing positively to the "
        "residence community through leadership, wellness, and cross-cultural collaboration."
    )

    # ---------- Education ----------
    add_section_heading(doc, "Education")
    add_bullet(doc, "  ·  Jimma University, Ethiopia  ·  2019",
               bold_lead="Bachelor of Science in Information Technology")
    add_bullet(doc, "  ·  Centennial College, Canada  ·  Expected Fall 2026",
               bold_lead="Postgraduate Certificate in Workplace Wellness and Health Promotion")

    # ---------- Certifications ----------
    add_section_heading(doc, "Certifications & Professional Development")
    add_bullet(doc, "  ·  AWS AI Practitioner Challenge", bold_lead="Cloud & AI —")
    add_bullet(doc, "  ·  Gemini in Google Drive", bold_lead="")
    add_bullet(doc, "  ·  Describe Cloud Computing", bold_lead="")
    add_bullet(doc, "  ·  Introduction to Critical Infrastructure Protection (OPSWAT)",
               bold_lead="Cybersecurity & Infrastructure —")
    add_bullet(doc, "  ·  Cyber Security Operations Job Simulation (Forage)", bold_lead="")
    add_bullet(doc, "  ·  Developing Emotional Intelligence (Saint Louis University)",
               bold_lead="Professional Skills —")
    add_bullet(doc, "  ·  Maze 101: Getting Started with Maze (Maze University)", bold_lead="")
    add_bullet(doc, "  ·  Hospitality Work Experience (Springpod / Marriott Bonvoy)", bold_lead="")

    # ---------- Professional Experience ----------
    add_section_heading(doc, "Professional Experience")
    add_job_title(doc, "IT Support Specialist",
                  "Star Development Association, Ethiopia    2022 – 2025")
    add_bullet(doc, "Resolved 95% of technical support tickets within SLA timelines, reducing operational downtime by 15%.")
    add_bullet(doc, "Mentored 10+ staff members on troubleshooting best practices, strengthening team capability and service consistency.")

    add_job_title(doc, "Data Entry & Database Support Officer",
                  "Dera City Administration, Ethiopia    2020 – 2022")
    add_bullet(doc, "Validated and maintained 50,000+ records with 99% accuracy across municipal databases.")
    add_bullet(doc, "Improved data quality controls, reducing duplicate records by 40% and supporting reliable reporting for city leadership.")

    # ---------- Key Projects ----------
    add_section_heading(doc, "Key Projects")
    add_bullet(doc, "Used Python and Power BI to improve IT support response efficiency by 20%.",
               bold_lead="IT Support Data Analysis —")
    add_bullet(doc, "Designed interactive Power BI dashboards to monitor operational performance.",
               bold_lead="Power BI Dashboarding —")
    add_bullet(doc, "Conducted vulnerability assessments and recommended mitigation strategies.",
               bold_lead="Cybersecurity Risk Simulation —")
    add_bullet(doc, "Implemented workflows using Google Drive and AI tools to enhance team efficiency.",
               bold_lead="Cloud Collaboration —")

    # ---------- Community & Leadership ----------
    add_section_heading(doc, "Community & Leadership")
    add_bullet(doc,
        " (2014–2019) — Provided digital literacy training to individuals with limited skills, building inclusive access to technology.",
        bold_lead="Volunteer Computer Tutor")
    add_bullet(doc,
        " Dedicated to future volunteerism and active contribution to the Centennial College residence community through peer mentorship, wellness initiatives, and cross-cultural engagement.",
        bold_lead="Commitment to Centennial:")

    # ---------- Awards & Interests ----------
    add_section_heading(doc, "Awards & Interests")
    add_bullet(doc, " Entrance Scholarship Recipient (Centennial College); Recognized for service excellence.",
               bold_lead="Awards —")
    add_bullet(doc, " AI-assisted analytics, digital transformation, wellness, and optimizing system performance through technical problem-solving.",
               bold_lead="Interests —")


def build_references(doc):
    """Page 3: References."""
    add_page_break(doc)
    add_centered_title(doc, "References", size=22)
    add_gold_rule(doc)
    add_centered_subtitle(doc, "Available upon request, or as listed below with permission.")
    doc.add_paragraph()

    add_reference_card(
        doc,
        name="Abdurezak Jeylan Abdulkadir",
        title="General Manager, Star Development Association (SDA), Dera, Ethiopia",
        relationship=(
            "Direct organizational head; signed my Experience Letter "
            "(Ref. SDA/EXPOS/25, dated 30/12/2025) confirming my role as "
            "IT Service Support Officer (2022–2025) and my contributions to "
            "IT support, systems, and project operations."
        ),
        contact_line="Email: Star3545@gmail.com  (SDA organizational)    |    Phone: +251 945 332 535  (SDA office)",
    )
    add_reference_card(
        doc,
        name="Saladin Musa",
        title="Lecturer and Researcher, Konbolcha Institute of Technology (former of Jimma University)",
        relationship=(
            "Academic reference; can speak to scholarly aptitude, training, "
            "and growth in Information Technology."
        ),
        contact_line="Email: Salahadin.seid@kiot.edu.et    |    Phone: +251 912 274 777    |    P.O. Box 378",
    )
    add_reference_card(
        doc,
        name="Kebebew Ababu, M.Sc.",
        title="Lecturer, Jimma Institute of Technology",
        relationship=(
            "Academic reference from the institution where I completed my "
            "B.Sc. in IT; can speak to academic performance and readiness "
            "for postgraduate study."
        ),
        contact_line="Email: Kebebew.ababu@ju.edu.et    |    Phone: +251 913 774 808    |    P.O. Box 378",
    )

    # Footnote
    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(16)
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note_p.add_run("Additional references and letters of recommendation available on request.")
    nr.italic = True
    nr.font.size = Pt(9.5)
    set_run_color(nr, MUTED)


def build_personal_statement(doc):
    """Pages 4-5: Personal Statement."""
    add_page_break(doc)
    add_centered_title(doc, "Personal Statement", size=22)
    add_centered_subtitle(doc, "Centennial College 2026 Residence Scholarship  ·  Bilal Abdulkadir Muhammed")
    add_gold_rule(doc)
    doc.add_paragraph()

    paragraphs = [
        "Five years ago, on a Tuesday afternoon in Dera, I sat beside a 58-year-old woman named "
        "Tigist who had been waiting three hours to renew her municipal ID. The database was "
        "down — again. As the Data Entry & Database Support Officer, I knew the problem was not "
        "the server — it was a workflow no one had ever questioned. I spent the next two weeks "
        "mapping the process, standardizing input fields, and training two colleagues. Within a "
        "month, duplicate records had fallen by 40%, and Tigist's daughter was able to register "
        "her newborn in under twenty minutes. That day taught me something I carry into every "
        "project I touch: systems are never just technical. They are human.",

        "That belief is what brings me to Centennial College. For six years I have worked at "
        "the intersection of technology and people — first as a database officer serving 50,000+ "
        "municipal records, then as an IT Support Specialist where I resolved 95% of tickets "
        "within SLA and mentored more than ten colleagues through their growth. The work was "
        "deeply satisfying, and it was not enough. Every ticket, every data field, every system "
        "I built was ultimately a question of wellbeing: was this person able to do their job "
        "without burning out? Was this family able to access the service they needed? Was this "
        "workplace actually well, or merely functional? I want to move closer to those questions. "
        "The Workplace Wellness and Health Promotion program at Centennial is the bridge I have "
        "been looking for — a chance to pair my analytical mindset and operational experience "
        "with formal training in mental health, occupational health, and the design of healthier "
        "workplaces.",

        "Relocating from Addis Ababa to Canada is not a small step, and I do not approach it "
        "lightly. I have spent the past year preparing — earning certifications in cloud "
        "computing (AWS, Microsoft), in cybersecurity infrastructure (OPSWAT, Forage), and "
        "intentionally, in developing emotional intelligence (Saint Louis University) and "
        "hospitality practice (Springpod / Marriott Bonvoy). Each one was a deliberate signal to "
        "myself that the next chapter requires more than technical skill. It requires the ability "
        "to read a room, to listen across cultures, and to lead with care. These are the muscles "
        "I want to strengthen at Centennial.",

        "If I am fortunate enough to receive this scholarship, I will bring that same "
        "orientation to the residence community. I have been a Volunteer Computer Tutor since "
        "2014, teaching digital literacy to neighbours and elders who had been left behind by "
        "the rapid spread of mobile technology. I know what it feels like to be the person in "
        "the room who does not yet speak the language of a new system, and I know the dignity "
        "of being helped without condescension. I want to be that person for incoming students "
        "at Centennial — someone who helps a first-generation international student navigate "
        "the student portal, who organizes a peer wellness check-in during exam season, or who "
        "simply listens when homesickness arrives unannounced.",

        "Centennial's mission is to cultivate graduates who are not only professionally capable "
        "but also community-minded. I want to be one of those graduates. I want to enter the "
        "workforce in Canada with the technical fluency I have built over six years, the "
        "wellness training Centennial will give me, and the lived understanding that every "
        "system — a database, a workplace, a residence hall — is, at its core, a collection of "
        "human beings trying to do well by one another. Thank you for considering my "
        "application. I will work hard to make this investment worth your faith.",
    ]
    for para in paragraphs:
        add_essay_paragraph(doc, para)

    # Signature block
    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(12)
    sig.paragraph_format.space_after = Pt(0)
    r = sig.add_run("Respectfully submitted,")
    r.italic = True
    r.font.size = Pt(10.5)
    set_run_color(r, MUTED)

    sig2 = doc.add_paragraph()
    sig2.paragraph_format.space_after = Pt(0)
    r = sig2.add_run("Bilal Abdulkadir Muhammed")
    r.bold = True
    r.font.size = Pt(12)
    set_run_color(r, ACCENT)

    sig3 = doc.add_paragraph()
    sig3.paragraph_format.space_after = Pt(0)
    r = sig3.add_run("Addis Ababa, Ethiopia  ·  June 2026")
    r.font.size = Pt(10)
    set_run_color(r, MUTED)


def build_leadership_essay(doc):
    """Pages 6-7: Leadership Essay."""
    add_page_break(doc)
    add_centered_title(doc, "Leadership Essay", size=22)
    add_centered_subtitle(doc, "Centennial College 2026 Residence Scholarship  ·  Bilal Abdulkadir Muhammed")
    add_gold_rule(doc)
    doc.add_paragraph()

    paragraphs = [
        "When I joined Dera City Administration as a Data Entry & Database Support Officer, my "
        "role was, on paper, narrow: enter records, verify them, keep the system running. In "
        "practice, I was standing at the front line of a city trying to serve tens of thousands "
        "of residents through a paper-and-spreadsheet system that was buckling under its own "
        "weight. My responsibilities grew to include high-volume record entry, ongoing data "
        "quality checks, and daily collaboration with administrative staff whose work depended "
        "entirely on the accuracy of what I and my colleagues put into the database. I came to "
        "understand, very quickly, that the smallest typo on my end could become a three-hour "
        "wait for a citizen at the front counter.",

        "I was appointed to the role on the strength of my IT background, but the leadership "
        "part of it was not assigned — it was earned. Within my first year, my supervisor began "
        "routing newer colleagues to me for guidance on the digital entry procedures I had helped "
        "refine. What started as informal shoulder-to-shoulder help became something more "
        "structured: I designed and led training sessions on the new digital registration forms, "
        "walked team members through verification protocols, and served as the bridge between "
        "the technical side of the system and the administrative staff who used it every day. "
        "The role had grown into a leadership position, and I had grown into it.",

        "The goal our team set was deceptively simple: move from manual, error-prone records to "
        "a reliable digital database that the city could actually build public services on. I "
        "led the effort to design standardized data entry protocols, implement new digital "
        "registration forms, and embed verification steps at the point of entry rather than "
        "after the fact. The results were concrete: a measurable reduction in entry errors, a "
        "streamlined workflow for the entire administrative team, and — most importantly — "
        "accurate citizen profiles that allowed public services to be delivered faster and more "
        "reliably. When a resident came in to renew an ID, register a child, or update a record, "
        "they were no longer held hostage by the inconsistency of our data. That shift mattered.",

        "It did not happen smoothly. The most difficult stretch came when we surfaced the scale "
        "of the problem we had actually inherited: fragmented, incomplete, and sometimes "
        "contradictory resident records that had been accumulating for years. The team was "
        "demoralized, and there was real resistance to the additional work a full audit would "
        "require. I responded by initiating a structured data audit, building a simple "
        "verification checklist that anyone on the team could use, and sitting with colleagues "
        "one-on-one to walk through the inconsistencies we were finding. The setback taught me "
        "that leadership is rarely about a clean plan — it is about doing the slow, unglamorous "
        "work of bringing people with you.",

        "Looking back, the most important lesson is that technical skill is necessary but "
        "never sufficient. Knowing how to design a database, or how to write a query, does not "
        "by itself change an organization. What changes an organization is the ability to "
        "communicate clearly, to listen to resistance rather than override it, and to build "
        "consensus around a better way of working. If I could lead this initiative again, I "
        "would have introduced collaborative training sessions from day one, rather than waiting "
        "until the new protocols were already drafted. Buy-in earlier would have saved us weeks "
        "of friction. I carry that lesson with me now, and I am eager to apply it at Centennial "
        "— whether that means helping residence students navigate a new wellness resource, "
        "supporting a peer through the transition to Canadian academic life, or simply being the "
        "person in the room who turns individual effort into shared progress.",
    ]
    for para in paragraphs:
        add_essay_paragraph(doc, para)


# ============================================================================
# PAGE SETUP
# ============================================================================
def setup_document(doc):
    """Configure page geometry, default style, header, and footer with page numbers."""
    # Letter, narrow margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    # Default body style: Calibri 10.5, 1.15 line spacing
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    # Running header (right-aligned, italic, small)
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Bilal Abdulkadir Muhammed  |  Resume")
    hr.font.size = Pt(9)
    hr.italic = True
    set_run_color(hr, MUTED)

    # Running footer: title + dynamic PAGE field
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fld_begin = _oxml("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = _oxml("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = _oxml("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = _oxml("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_num = _oxml("w:t")
    page_num.text = "1"

    prefix = fp.add_run("Centennial College 2026 Residence Scholarship Application  |  Page ")
    prefix.font.size = Pt(9)
    prefix.italic = True
    set_run_color(prefix, MUTED)

    page_run = fp.add_run()
    page_run.font.size = Pt(9)
    page_run.italic = True
    set_run_color(page_run, MUTED)
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_sep)
    page_run._r.append(page_num)
    page_run._r.append(fld_end)


# ============================================================================
# MAIN
# ============================================================================
def main():
    parser = argparse.ArgumentParser(
        description="Build Bilal's Centennial College 2026 Residence Scholarship DOCX package."
    )
    parser.add_argument(
        "--output", "-o",
        default=DEFAULT_OUTPUT,
        help=f"Output DOCX path (default: {DEFAULT_OUTPUT})",
    )
    args = parser.parse_args()

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    setup_document(doc)
    build_resume(doc)
    build_references(doc)
    build_personal_statement(doc)
    build_leadership_essay(doc)
    doc.save(output)

    print(f"Saved: {output}  ({output.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
